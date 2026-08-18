# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE = os.path.dirname(os.path.abspath(__file__))
SCREEN = os.path.join(BASE, "screenshots")
OUT = os.path.join(BASE, "Laporan Jejak Rimba v3.docx")

doc = Document()

# ---- Global styles ----
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
pf = normal.paragraph_format
pf.line_spacing = 1.5
pf.space_after = Pt(0)

for sname, size, bold in [("Heading 1", 14, True), ("Heading 2", 12, True)]:
    st = doc.styles[sname]
    st.font.name = "Times New Roman"
    st.font.size = Pt(size)
    st.font.bold = bold
    st.font.color.rgb = RGBColor(0, 0, 0)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

sec = doc.sections[0]
sec.top_margin = Cm(3)
sec.bottom_margin = Cm(3)
sec.left_margin = Cm(3)
sec.right_margin = Cm(4)


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)


add_page_number(sec)


def para(text, align=None, size=None, bold=False, italic=False, spacing=None,
         space_before=None, space_after=None, indent=True, first_line=Cm(1.25)):
    p = doc.add_paragraph()
    if align is None:
        align = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.alignment = align
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(size) if size else normal.font.size
    r.font.bold = bold
    r.font.italic = italic
    if spacing:
        p.paragraph_format.line_spacing = spacing
    if indent and align == WD_ALIGN_PARAGRAPH.JUSTIFY:
        p.paragraph_format.first_line_indent = first_line
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def heading(text, level=1):
    doc.add_heading(text, level=level)


def bab_heading(no, judul):
    p = doc.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(no)
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0, 0, 0)
    p2 = doc.add_paragraph(style="Heading 1")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(judul)
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(14)
    r2.font.bold = True
    r2.font.color.rgb = RGBColor(0, 0, 0)


def dot_list(entries, sub=False):
    for label, page, lvl in entries:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Cm(0.75 if (sub or lvl) else 0)
        add_tab_dot(p)
        r = p.add_run(label)
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
        tr = p.add_run()
        tr._r.append(OxmlElement("w:tab"))
        tr.font.name = "Times New Roman"
        tr.font.size = Pt(12)
        r2 = p.add_run(page)
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(12)


def numbered_item(num, text):
    para(f"{num}. {text}", spacing=1.5, indent=False)


def blank(n=1):
    for _ in range(n):
        doc.add_paragraph()


def add_image(path, width_cm, caption):
    doc.add_picture(path, width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    blank(1)
    para(caption, align=WD_ALIGN_PARAGRAPH.CENTER, size=11, indent=False)


def add_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(11)
        r.font.name = "Times New Roman"
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(11)
            r.font.name = "Times New Roman"
    return t


TOC_ENTRIES = [
    ("KATA PENGANTAR", "i", 0),
    ("DAFTAR ISI", "ii", 0),
    ("DAFTAR TABEL", "iii", 0),
    ("DAFTAR GAMBAR", "iii", 0),
    ("DAFTAR LAMPIRAN", "iii", 0),
    ("BAB 1 PENDAHULUAN", "1", 0),
    ("1.1 Latar Belakang", "1", 1),
    ("1.2 Tujuan", "2", 1),
    ("1.3 Manfaat", "2", 1),
    ("BAB 2 KAJIAN PUSTAKA", "3", 0),
    ("2.1 Website", "3", 1),
    ("2.2 Jejak Rimba", "3", 1),
    ("2.3 Next.js", "3", 1),
    ("2.4 Supabase", "4", 1),
    ("2.5 Midtrans", "4", 1),
    ("BAB 3 PROSES KERJA", "5", 0),
    ("3.1 Desain Produk", "5", 1),
    ("3.2 Alur Perencanaan dan Pelaksanaan", "5", 1),
    ("3.3 Alat dan Bahan", "6", 1),
    ("3.4 Jadwal Projek", "6", 1),
    ("BAB 4 HASIL PROYEK", "7", 0),
    ("4.1 Hasil", "7", 1),
    ("BAB 5 PENUTUP", "9", 0),
    ("5.1 Kesimpulan", "9", 1),
    ("5.2 Saran", "9", 1),
    ("DAFTAR RUJUKAN", "10", 0),
    ("PROFIL PENULIS", "11", 0),
]


def add_tab_dot(p):
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:leader"), "dot")
    tab.set(qn("w:pos"), "7717")
    tabs.append(tab)
    pPr.append(tabs)


def add_toc_field():
    def run_of(p, text):
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
        return r

    # Paragraf pembuka field: begin + instr + separate
    p0 = doc.add_paragraph()
    r0 = p0.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
    it.text = 'TOC \\o "1-2" \\h \\z \\u'
    fs = OxmlElement("w:fldChar"); fs.set(qn("w:fldCharType"), "separate")
    r0._r.append(f1); r0._r.append(it); r0._r.append(fs)

    # Isi daftar isi (cached result) supaya langsung tampil
    for label, page, lvl in TOC_ENTRIES:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Cm(0.75 if lvl else 0)
        add_tab_dot(p)
        run_of(p, label)
        tab_run = p.add_run()
        tab_run._r.append(OxmlElement("w:tab"))
        tab_run.font.name = "Times New Roman"
        tab_run.font.size = Pt(12)
        run_of(p, page)

    # Paragraf penutup field: end
    pz = doc.add_paragraph()
    rz = pz.add_run()
    fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end")
    rz._r.append(fe)


# Update fields on open (biar TOC auto tampil)
settings = doc.settings.element
updateFields = OxmlElement("w:updateFields")
updateFields.set(qn("w:val"), "true")
settings.append(updateFields)


# ============ COVER ============
para("LAPORAN PROJECT BASED LEARNING", align=WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True)
para("WEBSITE JEJAK RIMBA", align=WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True)
para("KONSENTRASI KEAHLIAN REKAYASA PERANGKAT LUNAK", align=WD_ALIGN_PARAGRAPH.CENTER, size=13, bold=True)
para("SEMESTER GANJIL 2026/2027", align=WD_ALIGN_PARAGRAPH.CENTER, size=13, bold=True)
blank(4)
para("PENYUSUN :", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True)
para("[Nama Lengkap Penulis]", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
para("NIS [Nomor Induk Siswa]", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
blank(2)
para("PEMBIMBING :", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True)
para("[Nama Guru Pembimbing, M.Pd]", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
blank(4)
para("SMK NEGERI 4 MALANG", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True)
para("Jalan Tanimbar Nomor 22 Malang, Jawa Timur 65117", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
para("Telepon (0341) 353798", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
para("Laman www.smkn4malang.sch.id, Pos-el mail@smkn4malang.sch.id", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
doc.add_page_break()

# ============ KATA PENGANTAR ============
heading("KATA PENGANTAR", 1)
para("Assalamu'alaikum Wr. Wb.")
para("Puji syukur kehadirat Allah SWT yang telah melimpahkan rahmat, taufik, dan hidayah-Nya sehingga penulis dapat menyelesaikan laporan Project Based Learning dengan judul \"Website Jejak Rimba\" ini dengan baik. Shalawat serta salam senantiasa tercurah kepada junjungan kita Nabi Muhammad SAW beserta keluarga, sahabat, dan seluruh umatnya hingga akhir zaman. Penyusunan laporan ini merupakan salah satu bentuk pertanggungjawaban penulis atas pelaksanaan pembelajaran berbasis proyek selama satu semester berjalan.")
para("Penulis menyadari sepenuhnya bahwa laporan ini tidak akan terselesaikan tanpa bantuan, bimbingan, serta dukungan dari berbagai pihak. Oleh karena itu, penulis ingin menyampaikan rasa hormat dan ucapan terima kasih yang tulus kepada:")
numbered_item(1, "Bapak/Ibu Kepala SMK Negeri 4 Malang yang telah memberikan izin, kesempatan, serta dukungan kepada penulis dalam pelaksanaan kegiatan Project Based Learning (PjBL) sehingga penyusunan laporan ini dapat terlaksana dengan baik.")
numbered_item(2, "Bapak/Ibu Ketua Program Keahlian Rekayasa Perangkat Lunak (RPL) SMK Negeri 4 Malang yang telah memberikan arahan, dukungan, serta kerja sama selama pelaksanaan kegiatan Project Based Learning (PjBL) dan penyusunan laporan ini.")
numbered_item(3, "Bapak/Ibu Guru Pembimbing di SMK Negeri 4 Malang yang telah memberikan bimbingan, masukan, serta pendampingan kepada penulis selama pelaksanaan kegiatan Project Based Learning (PjBL) sehingga laporan ini dapat diselesaikan dengan baik.")
numbered_item(4, "Bapak/Ibu Wali Kelas yang telah memberikan arahan, dukungan, serta membantu mengoordinasikan peserta didik selama pelaksanaan kegiatan Project Based Learning (PjBL), sehingga kegiatan dapat berlangsung dengan baik dan laporan ini dapat diselesaikan.")
numbered_item(5, "Kedua orang tua penulis yang senantiasa memberikan doa, kasih sayang, serta dukungan moral dan material selama penulis menempuh pendidikan dan menyelesaikan laporan ini.")
numbered_item(6, "Teman-teman yang telah memberikan dukungan, motivasi, bantuan, serta berbagi pengalaman dan masukan selama pelaksanaan kegiatan Project Based Learning (PjBL) hingga penyusunan laporan ini dapat diselesaikan.")
para("Penulis menyadari bahwa laporan ini masih jauh dari kata sempurna. Kritik dan saran yang bersifat membangun sangat penulis harapkan demi perbaikan di masa yang akan datang. Akhirnya, penulis berharap laporan ini dapat memberikan manfaat bagi siapa saja yang membacanya, khususnya bagi peserta didik Program Keahlian Rekayasa Perangkat Lunak.")
para("Wassalamu'alaikum Wr. Wb.")
blank(2)
para("Malang, [Tanggal] 2026", align=WD_ALIGN_PARAGRAPH.RIGHT)
para("Penulis,", align=WD_ALIGN_PARAGRAPH.RIGHT)
blank(3)
para("[Nama Lengkap Penulis]", align=WD_ALIGN_PARAGRAPH.RIGHT)
doc.add_page_break()

# ============ DAFTAR ISI (otomatis) ============
heading("DAFTAR ISI", 1)
add_toc_field()
doc.add_page_break()

# ============ DAFTAR TABEL & GAMBAR ============
heading("DAFTAR TABEL", 1)
dot_list([
    ("Tabel 3.1 Jadwal Projek", "6", 0),
    ("Tabel 4.1 Data Alat pada Katalog Jejak Rimba", "8", 0),
])
doc.add_page_break()
heading("DAFTAR GAMBAR", 1)
dot_list([
    ("Gambar 3.1 Wireframe Halaman Masuk", "5", 0),
    ("Gambar 3.2 Alur Perencanaan dan Pelaksanaan", "5", 0),
    ("Gambar 3.3 Diagram ERD Jejak Rimba", "6", 0),
    ("Gambar 3.4 Use Case Diagram Jejak Rimba", "6", 0),
    ("Gambar 3.5 Activity Diagram Pemesanan Alat", "6", 0),
    ("Gambar 4.1 Tampilan Halaman Utama", "8", 0),
    ("Gambar 4.2 Tampilan Halaman Detail Alat", "8", 0),
])
doc.add_page_break()
heading("DAFTAR LAMPIRAN", 1)
dot_list([
    ("Lampiran 1 Struktur Tabel Basis Data", "12", 0),
    ("Lampiran 2 Foto Kegiatan", "14", 0),
])
doc.add_page_break()

# ============ BAB 1 ============
bab_heading("BAB 1", "PENDAHULUAN")
heading("1.1 Latar Belakang", 2)
para("Kegiatan mendaki gunung dan berkemah mengalami peningkatan yang cukup signifikan dalam beberapa tahun terakhir. Perkembangan ini tidak lepas dari maraknya konten petualangan yang dibagikan di berbagai platform media sosial seperti Instagram, TikTok, dan Twitter. Banyak kreator konten serta influencer yang rutin menceritakan pengalamannya berada di alam terbuka, mulai dari perjalanan menuju puncak, mendirikan tenda di tengah hutan, hingga menikmati panorama matahari terbit. Paparan konten semacam ini secara tidak langsung menumbuhkan ketertarikan masyarakat, terutama generasi muda, untuk ikut merasakan pengalaman serupa.")
para("Namun, keinginan untuk mencoba kegiatan tersebut sering kali terbentur oleh satu kendala yang cukup besar, yaitu keterbatasan peralatan. Perlengkapan pendakian seperti tenda, carrier, sleeping bag, matras, kompor, dan jaket gunung bukanlah barang murah. Harga satu unit tenda berkualitas saja dapat mencapai ratusan ribu rupiah, belum lagi peralatan pendukung lainnya. Bagi masyarakat umum yang hanya berencana mendaki sekali atau dua kali dalam setahun, membeli seluruh perlengkapan tersebut tentu tidak efisien. Alat yang dibeli mahal justru hanya tersimpan di rumah setelah kegiatan selesai.")
para("Di sisi lain, keberadaan penyedia jasa sewa alat di wilayah Malang Raya belum diimbangi dengan sistem layanan yang baik. Banyak penyedia yang masih mengandalkan cara manual, seperti menunggu pesan melalui WhatsApp atau mengunggah ketersediaan alat di media sosial. Cara ini membuat calon penyewa harus berkomunikasi satu per satu dengan setiap penyedia hanya untuk mengetahui apakah alat yang dibutuhkan masih tersedia. Akibatnya, proses pencarian menjadi lambat, tidak terstruktur, dan sering kali menimbulkan kekecewaan karena alat sudah disewa orang lain atau harga tidak sesuai dengan yang tertera sebelumnya.")
para("Persoalan tersebut juga berdampak pada penyedia itu sendiri. Jangkauan pemasaran yang terbatas pada lingkaran pertemanan terdekat membuat alat-alat yang dimiliki tidak selalu tersewa, sehingga potensi pendapatan menjadi tidak maksimal. Padahal, jika dikelola dengan baik, usaha sewa alat outdoor memiliki pasar yang cukup menjanjikan, terlebih Malang dikenal sebagai salah satu daerah dengan banyak destinasi pendakian dan perkemahan.")
para("Berdasarkan permasalahan tersebut, penulis berinisiatif membangun sebuah platform bernama Jejak Rimba. Platform ini berbentuk website yang mempertemukan penyewa dengan penyedia alat camping dan mendaki secara daring. Melalui website ini, pengguna dapat menjelajahi katalog alat, melihat jumlah stok secara langsung, membandingkan harga dan lokasi penyedia, serta melakukan pemesanan dan pembayaran tanpa perlu bertemu langsung. Dengan demikian, proses sewa alat yang sebelumnya rumit dapat disederhanakan menjadi beberapa langkah yang mudah dan transparan.")
heading("1.2 Tujuan", 2)
para("Tujuan pengembangan website Jejak Rimba adalah sebagai berikut.")
numbered_item(1, "Menyediakan platform persewaan alat mendaki dan berkemah yang dapat diakses kapan saja selama 24 jam, sehingga pengguna tidak lagi bergantung pada jam operasional penyedia.")
numbered_item(2, "Memberikan pengalaman transaksi yang aman, transparan, dan efisien, baik bagi pengguna maupun penyedia alat melalui sistem pemesanan dan pembayaran daring.")
numbered_item(3, "Membangun ekosistem yang saling menguntungkan antara pengguna dan penyedia alat persewaan, khususnya di wilayah Malang Raya.")
heading("1.3 Manfaat", 2)
para("Manfaat yang diharapkan dari pengembangan website Jejak Rimba antara lain:")
numbered_item(1, "Bagi pengguna, mempermudah proses pencarian dan penyewaan alat camping yang tersedia tanpa harus menghubungi penyedia satu per satu, serta memperoleh informasi stok, harga, dan lokasi secara real time.")
numbered_item(2, "Bagi penyedia, memperluas jangkauan pasar sehingga alat yang dimiliki dapat tersewa lebih banyak dan pendapatan usaha menjadi lebih optimal.")
numbered_item(3, "Bagi penulis, menambah wawasan serta pengalaman praktis dalam merancang dan membangun aplikasi berbasis website menggunakan teknologi modern seperti Next.js, Supabase, dan Midtrans.")
doc.add_page_break()

# ============ BAB 2 ============
bab_heading("BAB 2", "KAJIAN PUSTAKA")
heading("2.1 Website", 2)
para("Website adalah kumpulan halaman yang saling terhubung dan dapat diakses melalui internet menggunakan peramban atau browser. Setiap website memiliki identitas berupa alamat domain yang unik, sehingga dapat ditemukan oleh pengguna di seluruh dunia. Halaman-halaman yang terdapat di dalamnya dapat memuat berbagai jenis konten, mulai dari teks, gambar, video, hingga formulir interaktif yang memungkinkan pengguna berinteraksi dengan sistem.")
para("Secara garis besar, website dapat dikelompokkan menjadi dua jenis berdasarkan cara kerjanya. Jenis pertama adalah website statis, yaitu website yang menampilkan konten yang tetap dan tidak berubah kecuali file sumbernya diubah secara manual. Jenis kedua adalah website dinamis, yaitu website yang kontennya dapat berubah menyesuaikan interaksi pengguna dan data yang tersimpan di dalam basis data. Jejak Rimba termasuk ke dalam website dinamis karena katalog alat, status stok, dan data pemesanan diambil secara langsung dari basis data setiap kali halaman diminta.")
heading("2.2 Jejak Rimba", 2)
para("Jejak Rimba adalah platform persewaan alat camping dan mendaki yang beroperasi di wilayah Malang Raya, mencakup Malang Kota, Batu, dan Lawang. Melalui website ini, pengguna dapat menjelajahi katalog alat yang tersedia, melihat ketersediaan stok secara langsung, membandingkan harga serta lokasi penyedia, kemudian melakukan pemesanan dan pembayaran secara daring. Seluruh proses, dari pemilihan alat hingga konfirmasi pembayaran, dapat diselesaikan dalam satu alur yang terstruktur.")
para("Platform ini mempertemukan dua pihak yang saling membutuhkan. Pihak pertama adalah penyewa, yaitu masyarakat umum yang membutuhkan alat untuk kegiatan berkemah maupun mendaki. Pihak kedua adalah penyedia, yaitu pemilik usaha rental yang menyewakan perlengkapan outdoor. Dengan adanya platform, interaksi di antara keduanya difasilitasi secara digital sehingga proses komunikasi, pencatatan pemesanan, dan pembayaran menjadi lebih cepat, tertib, dan dapat dipertanggungjawabkan.")
heading("2.3 Next.js", 2)
para("Next.js adalah kerangka kerja atau framework yang dibangun di atas React untuk mengembangkan aplikasi website. Next.js mendukung berbagai teknik rendering seperti server-side rendering dan static generation, sehingga halaman dapat dimuat dengan lebih cepat dan ramah terhadap mesin pencari. Framework ini juga menyediakan sistem routing berbasis folder, API routes, serta server actions untuk menangani logika bisnis yang berjalan di sisi server.")
para("Pada pengembangan Jejak Rimba, Next.js digunakan bersama dengan TypeScript sebagai bahasa pemrograman utama. TypeScript menambahkan pemeriksaan tipe data sehingga kode menjadi lebih aman, mudah dipelihara, dan terhindar dari berbagai kesalahan yang baru terdeteksi saat aplikasi berjalan. Untuk tampilan antarmuka, digunakan React sebagai pustaka komponen, Tailwind CSS untuk pengaturan gaya, serta Framer Motion untuk memberikan animasi yang halus dan menarik bagi pengguna.")
heading("2.4 Supabase", 2)
para("Supabase adalah layanan backend yang bersifat open source dan menyediakan berbagai kebutuhan pengembangan aplikasi, di antaranya basis data PostgreSQL, autentikasi pengguna, serta penyimpanan file. Salah satu keunggulan utama Supabase adalah dukungan terhadap Row Level Security, yaitu mekanisme keamanan yang diterapkan pada tingkat baris data. Dengan mekanisme ini, setiap pengguna hanya dapat mengakses data yang memang menjadi haknya, sehingga data pengguna lain tetap terlindungi.")
para("Autentikasi pada Jejak Rimba ditangani sepenuhnya oleh Supabase Auth. Pengguna dapat mendaftar dan masuk menggunakan alamat email beserta kata sandi, atau melalui akun Google. Sesi pengguna dikelola menggunakan cookie yang disegarkan secara otomatis agar tidak mudah kedaluwarsa. Seluruh data yang berkaitan dengan profil, alat, pemesanan, dan ulasan disimpan pada basis data PostgreSQL yang disediakan oleh Supabase.")
heading("2.5 Midtrans", 2)
para("Midtrans adalah penyedia layanan payment gateway yang banyak digunakan di Indonesia. Midtrans mendukung beragam metode pembayaran, mulai dari transfer bank, kartu kredit, hingga dompet digital. Dengan menggunakan layanan ini, aplikasi dapat menerima pembayaran dari berbagai bank dan metode tanpa perlu membangun koneksi langsung dengan sistem perbankan, sehingga proses integrasi menjadi lebih sederhana dan aman.")
para("Pada Jejak Rimba, Midtrans berperan dalam menangani proses pembayaran pesanan. Ketika pengguna melakukan pembayaran, status transaksi akan diperbarui secara otomatis melalui mekanisme webhook yang dikirim oleh Midtrans ke server aplikasi. Dengan cara ini, status pembayaran selalu mencerminkan kondisi yang sebenarnya tanpa bergantung pada tindakan pengguna, sehingga risiko kesalahan pencatatan dapat diminimalkan.")
doc.add_page_break()

# ============ BAB 3 ============
bab_heading("BAB 3", "PROSES KERJA")
heading("3.1 Desain Produk", 2)
para("Tahap desain dimulai dengan pembuatan wireframe untuk halaman-halaman utama. Wireframe merupakan gambaran sederhana yang menunjukkan tata letak elemen pada suatu halaman sebelum diubah menjadi kode program. Pembuatan wireframe bertujuan untuk memastikan bahwa alur dan struktur halaman sudah jelas, sehingga memudahkan proses pengembangan berikutnya. Dua wireframe yang dibuat lebih dahulu adalah halaman masuk dan halaman daftar.")
para("Halaman masuk dirancang dengan tombol kembali di pojok kiri atas, logo Jejak Rimba di bagian tengah, judul \"Masuk ke Akun Anda\", dua kolom input untuk email dan kata sandi, tautan lupa kata sandi, tombol masuk utama dengan warna aksen, pemisah \"atau\", tombol masuk menggunakan akun Google, serta tautan untuk pindah ke halaman daftar. Sementara itu, halaman daftar memiliki tombol kembali, judul \"Buat Akun Baru\", empat kolom input untuk nama lengkap, email, kata sandi, dan konfirmasi kata sandi, tombol daftar utama, pemisah \"atau\", tombol daftar dengan Google, serta tautan untuk pindah ke halaman masuk.")
para("Selain dua halaman tersebut, penulis juga merancang tampilan katalog dan halaman detail alat. Desain katalog menampilkan daftar alat dalam bentuk kartu yang memuat foto, nama, kategori, harga per hari, lokasi, dan indikator ketersediaan stok. Halaman detail alat menampilkan informasi yang lebih lengkap, termasuk kondisi alat dan nama penyedia, serta tombol untuk melakukan pemesanan.")
para("[Gambar 3.1 Wireframe Halaman Masuk]", align=WD_ALIGN_PARAGRAPH.CENTER)
heading("3.2 Alur Perencanaan dan Pelaksanaan", 2)
para("Pengerjaan Jejak Rimba disusun melalui beberapa tahap yang dilakukan secara berurutan. Penyusunan tahapan ini dimaksudkan agar setiap pekerjaan dapat diselesaikan secara tuntas sebelum melanjutkan ke tahap berikutnya, sehingga hasil akhirnya sesuai dengan rencana.")
numbered_item(1, "Analisis kebutuhan. Tahap ini dimulai dengan mengidentifikasi fitur yang dibutuhkan oleh pengguna dan penyedia. Fitur utama yang ditetapkan meliputi katalog alat, sistem pemesanan, dan pembayaran daring. Penulis juga menelaah kelemahan sistem sewa konvensional agar dapat diatasi pada perancangan.")
numbered_item(2, "Perancangan data. Penulis membuat diagram ERD untuk menggambarkan struktur data. Diagram ini menjadi dasar dalam menentukan delapan tabel utama beserta relasi antar tabel, sehingga struktur basis data dapat dirancang secara rapi dan mudah dikembangkan.")
numbered_item(3, "Implementasi basis data. Seluruh tabel yang telah dirancang dibuat pada Supabase. Pada tahap ini juga diterapkan aturan keamanan baris agar setiap pengguna hanya dapat mengakses data miliknya.")
numbered_item(4, "Pengembangan halaman. Halaman utama, katalog, detail alat, dan alur pemesanan dibangun menggunakan Next.js. Tampilan antarmuka disusun menggunakan Tailwind CSS agar responsif pada berbagai ukuran layar.")
numbered_item(5, "Penerapan autentikasi. Sistem masuk dan daftar pengguna dihubungkan dengan Supabase Auth. Selanjutnya, route yang bersifat pribadi seperti profil dan pemesanan dilindungi oleh middleware agar hanya dapat diakses oleh pengguna yang telah masuk.")
numbered_item(6, "Pengujian. Seluruh fitur diuji secara menyeluruh untuk memastikan alur pemesanan, pembayaran, dan pembaruan status berjalan dengan benar tanpa ditemui kendala yang berarti.")
para("[Gambar 3.2 Alur Perencanaan dan Pelaksanaan]", align=WD_ALIGN_PARAGRAPH.CENTER)
add_image(os.path.join(SCREEN, "diagram-erd.png"), 14.5, "Gambar 3.3 Diagram ERD Jejak Rimba")
add_image(os.path.join(os.path.expanduser("~"), "Downloads", "jejak_rimba_use_case_diagram.png"), 14.5, "Gambar 3.4 Use Case Diagram Jejak Rimba")
add_image(os.path.join(os.path.expanduser("~"), "Downloads", "jejak_rimba_activity_booking.png"), 14.5, "Gambar 3.5 Activity Diagram Pemesanan Alat")
heading("3.3 Alat dan Bahan", 2)
para("Dalam pengerjaan proyek ini, penulis memerlukan sejumlah perangkat keras dan perangkat lunak. Pemilihan perangkat dilakukan dengan mempertimbangkan kebutuhan pengembangan serta kemampuan yang dimiliki. Perangkat dan perangkat lunak tersebut di antaranya sebagai berikut.")
numbered_item(1, "Laptop: perangkat keras utama yang digunakan untuk menjalankan seluruh aplikasi yang dibutuhkan selama proses pengembangan.")
numbered_item(2, "Visual Studio Code: editor kode yang digunakan untuk menulis dan mengelola seluruh kode program.")
numbered_item(3, "Next.js dan TypeScript: kerangka kerja serta bahasa pemrograman yang digunakan untuk membangun aplikasi.")
numbered_item(4, "Tailwind CSS dan Framer Motion: digunakan untuk menata tampilan antarmuka dan menambahkan animasi.")
numbered_item(5, "Supabase: layanan yang digunakan sebagai basis data dan sistem autentikasi pengguna.")
numbered_item(6, "Midtrans: penyedia layanan pembayaran daring yang diintegrasikan ke dalam sistem.")
numbered_item(7, "Git dan GitHub: digunakan untuk mengelola versi kode dan menyimpan repository agar pekerjaan terdokumentasi.")
numbered_item(8, "Figma: digunakan dalam perancangan desain antarmuka sebelum diterjemahkan menjadi kode program.")
heading("3.4 Jadwal Projek", 2)
para("Pelaksanaan proyek disusun dalam bentuk jadwal agar setiap tahap pengerjaan dapat diselesaikan tepat waktu. Jadwal tersebut mencakup tiga tahap utama, yaitu perencanaan, pelaksanaan, dan pelaporan. Rincian jadwal dapat dilihat pada tabel berikut.")
add_table(["Tahapan", "Waktu", "Kegiatan"],
          [["Perencanaan", "14-16 April 2026", "Analisis kebutuhan sistem dan perancangan data"],
           ["", "17 April 2026", "Perancangan basis data (ERD)"],
           ["Pelaksanaan", "21 April - 16 Mei 2026", "Pembuatan basis data dan koneksi (Supabase)"],
           ["", "18 Mei 2026", "Implementasi halaman dinamis dan autentikasi"],
           ["", "19-20 Mei 2026", "Pembuatan katalog alat dan dashboard penyedia"],
           ["", "21 Mei 2026", "Implementasi alur pemesanan dan pembayaran"],
           ["", "22 Mei 2026", "Pengujian sistem"],
           ["Pelaporan", "4 Juni 2026", "Finalisasi laporan"],
           ["", "5 Juni 2026", "Presentasi akhir"]])
para("Tabel 3.1 Jadwal Projek", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
doc.add_page_break()

# ============ BAB 4 ============
bab_heading("BAB 4", "HASIL PROYEK")
heading("4.1 Hasil", 2)
para("Website Jejak Rimba berhasil dibangun sesuai dengan perancangan yang telah ditetapkan sebelumnya. Seluruh halaman utama yang direncanakan telah terwujud dan dapat berjalan dengan baik. Fitur-fitur tersebut disusun dengan memperhatikan kemudahan penggunaan agar seluruh kalangan, termasuk pengguna yang baru pertama kali mendaki, dapat mengoperasikannya tanpa kesulitan.")
numbered_item(1, "Halaman utama yang menampilkan sambutan, alur cara sewa, katalog alat unggulan, serta testimoni pengguna.")
add_image(os.path.join(SCREEN, "home.png"), 14.5, "Gambar 4.1 Tampilan Halaman Utama")
numbered_item(2, "Halaman katalog yang memuat daftar alat lengkap dengan filter kategori dan lokasi, sehingga pengguna dapat menyaring alat sesuai kebutuhan.")
para("[Screenshot halaman katalog]", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
numbered_item(3, "Halaman detail alat yang menampilkan harga per hari, stok, kondisi, lokasi, dan nama penyedia.")
add_image(os.path.join(SCREEN, "detail.png"), 14.5, "Gambar 4.2 Tampilan Halaman Detail Alat")
numbered_item(4, "Alur pemesanan yang terdiri dari pemilihan tanggal, pengisian data, dan konfirmasi pesanan.")
para("[Screenshot halaman pemilihan tanggal]", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
numbered_item(5, "Halaman masuk, daftar, dan lupa kata sandi yang terhubung dengan Supabase Auth.")
para("[Screenshot halaman masuk dan daftar]", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
numbered_item(6, "Halaman profil untuk mengelola data diri dan melihat riwayat pemesanan.")
para("[Screenshot halaman profil]", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
numbered_item(7, "Dashboard khusus untuk penyedia dalam mengelola alat dan melihat pesanan masuk.")
para("[Screenshot dashboard penyedia]", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
numbered_item(8, "Halaman informasi berupa FAQ, kebijakan privasi, serta syarat dan ketentuan.")
para("[Screenshot halaman FAQ]", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
para("Data alat yang ditampilkan pada katalog mengambil contoh nyata dari beberapa penyedia yang beroperasi di Malang Raya. Contoh data ini bertujuan untuk menggambarkan bagaimana informasi alat, mulai dari harga per hari hingga lokasi pengambilan, disajikan kepada pengguna. Beberapa contoh data alat dapat dilihat pada tabel berikut.")
add_table(["Nama Alat", "Kategori", "Harga/Hari", "Stok", "Lokasi"],
          [["Tenda Dome Consina 4P", "Tenda", "Rp35.000", "6", "Malang Kota"],
           ["Tenda Ultralight Eiger 2P", "Tenda", "Rp45.000", "3", "Batu"],
           ["Carrier Avtech 60L", "Carrier", "Rp20.000", "10", "Malang Kota"],
           ["Carrier Deuter 80L", "Carrier", "Rp28.000", "4", "Lawang"],
           ["Sleeping Bag Naturehike M400", "Sleeping Bag", "Rp15.000", "12", "Malang Kota"],
           ["Kompor Portable + Gas Windproof", "Kompor", "Rp12.000", "8", "Batu"],
           ["Matras Lipat Aluminium Foil", "Matras", "Rp8.000", "15", "Lawang"],
           ["Jaket Gunung Waterproof Eiger", "Jaket", "Rp25.000", "0", "Malang Kota"]])
para("Tabel 4.1 Data Alat pada Katalog Jejak Rimba", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
para("Selain katalog, hasil proyek ini juga mencakup sistem autentikasi dan alur pemesanan yang terintegrasi dengan pembayaran. Sistem ini memungkinkan pengguna untuk melihat status pesanan dari mulai menunggu konfirmasi hingga selesai. Integrasi dengan Midtrans memastikan bahwa setiap perubahan status pembayaran dapat diperbarui secara otomatis tanpa perlu campur tangan pengguna.")
doc.add_page_break()

# ============ BAB 5 ============
bab_heading("BAB 5", "PENUTUP")
heading("5.1 Kesimpulan", 2)
para("Berdasarkan hasil pengerjaan proyek ini, penulis dapat menarik beberapa kesimpulan. Pertama, website Jejak Rimba berhasil dibangun sebagai platform persewaan alat camping dan mendaki yang dapat diakses melalui peramban, baik dari perangkat desktop maupun ponsel. Seluruh halaman utama, mulai dari katalog hingga alur pemesanan, telah berjalan sesuai dengan perancangan awal.")
para("Kedua, fitur-fitur utama seperti katalog alat, sistem pemesanan, dan pembayaran daring terbukti mampu menjawab permasalahan yang diangkat pada latar belakang. Teknologi yang digunakan, yaitu Next.js untuk antarmuka serta Supabase untuk basis data dan autentikasi, memberikan fondasi yang kuat dan mudah dikembangkan. Melalui platform ini, calon pendaki dapat menemukan alat yang tersedia tanpa harus menghubungi penyedia satu per satu, sehingga prosesnya menjadi lebih cepat dan efisien.")
heading("5.2 Saran", 2)
para("Untuk pengembangan selanjutnya, penulis menyarankan beberapa perbaikan. Katalog dapat dilengkapi dengan fitur pencarian dan filter berdasarkan harga, ketersediaan tanggal, serta peringkat penyedia. Penambahan fitur peta interaktif juga akan membantu pengguna melihat lokasi pengambilan alat secara visual dan memilih titik yang paling dekat.")
para("Selain itu, sistem pengingat untuk pengembalian alat sebaiknya ditambahkan agar penyewa tidak melewatkan batas waktu yang telah disepakati. Penyedia juga akan terbantu apabila tersedia laporan transaksi yang dapat diunduh dalam bentuk berkas. Dengan berbagai penyempurnaan tersebut, pengalaman seluruh pihak yang terlibat dalam platform diharapkan dapat menjadi semakin nyaman dan bermanfaat.")
doc.add_page_break()

# ============ DAFTAR PUSTAKA ============
heading("DAFTAR PUSTAKA", 1)
refs = [
    "Framer. \"Framer Motion Documentation.\" Accessed [Tanggal], 2026. https://motion.dev.",
    "Kemdikbud. \"Kamus Besar Bahasa Indonesia.\" Accessed [Tanggal], 2026. https://kbbi.kemdikbud.go.id.",
    "Midtrans. \"Midtrans Documentation: Payment Gateway Indonesia.\" Accessed [Tanggal], 2026. https://docs.midtrans.com.",
    "Supabase. \"Supabase Documentation: Auth, Database, Row Level Security.\" Accessed [Tanggal], 2026. https://supabase.com/docs.",
    "Tailwind Labs. \"Tailwind CSS Documentation.\" Accessed [Tanggal], 2026. https://tailwindcss.com/docs.",
    "Vercel. \"Next.js Documentation.\" Accessed [Tanggal], 2026. https://nextjs.org/docs.",
]
for r in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(-1.25)
    run = p.add_run(r)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
doc.add_page_break()

# ============ PROFIL PENULIS ============
heading("PROFIL PENULIS", 1)
para("Penulis bernama [Nama Lengkap Penulis], lahir di [Tempat Lahir] pada tanggal [Tanggal Lahir]. Saat ini penulis merupakan peserta didik kelas XI Program Keahlian Rekayasa Perangkat Lunak (RPL) di SMK Negeri 4 Malang. Penulis memiliki minat yang kuat di bidang pengembangan perangkat lunak, khususnya dalam pembuatan website, desain antarmuka (UI/UX), serta pengelolaan basis data.")
para("Selama menempuh pendidikan di SMK Negeri 4 Malang, penulis aktif mengikuti berbagai kegiatan pembelajaran berbasis proyek yang bertujuan meningkatkan kompetensi di bidang teknologi informasi. Salah satu proyek yang pernah dikerjakan adalah website Jejak Rimba, yaitu platform persewaan alat camping dan mendaki. Melalui proyek tersebut, penulis memperoleh pengalaman berharga dalam menganalisis kebutuhan pengguna, merancang sistem, mengembangkan aplikasi, serta bekerja sama dengan rekan satu tim.")
para("Penulis berharap ilmu dan pengalaman yang diperoleh selama masa pendidikan dapat menjadi bekal untuk melanjutkan studi maupun berkarier di bidang teknologi informasi. Penulis juga ingin memberikan kontribusi positif bagi masyarakat melalui pengembangan perangkat lunak yang bermanfaat dan dapat digunakan secara luas.")
doc.add_page_break()

# ============ LAMPIRAN ============
heading("LAMPIRAN 1: STRUKTUR TABEL BASIS DATA", 1)
para("Tabel User (Profiles)", bold=True)
add_table(["No", "Nama Field", "Tipe", "Panjang", "Keterangan"],
          [["1", "id", "UUID", "-", "Primary Key"],
           ["2", "full_name", "text", "-", "Nama lengkap pengguna"],
           ["3", "email", "text", "-", "Alamat email"],
           ["4", "phone_number", "text", "-", "Nomor telepon"],
           ["5", "role", "enum", "-", "penyewa, vendor, admin"],
           ["6", "avatar_url", "text", "-", "Foto profil"]])
blank()
para("Tabel Equipment", bold=True)
add_table(["No", "Nama Field", "Tipe", "Panjang", "Keterangan"],
          [["1", "id", "UUID", "-", "Primary Key"],
           ["2", "vendor_id", "UUID", "-", "Foreign Key ke vendors"],
           ["3", "name", "text", "-", "Nama alat"],
           ["4", "category", "text", "-", "Kategori alat"],
           ["5", "price_per_day", "number", "-", "Harga per hari"],
           ["6", "total_stock", "integer", "-", "Jumlah stok"],
           ["7", "condition", "enum", "-", "Kondisi alat"]])
doc.add_page_break()
heading("LAMPIRAN 2: FOTO KEGIATAN", 1)
para("[Tempat foto kegiatan selama proses pengerjaan proyek]")

doc.save(OUT)
print("OK ->", OUT)
