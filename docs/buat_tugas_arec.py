# -*- coding: utf-8 -*-
# Tugas terpisah: Copywriting Landing Page (Matriks ARE-C)
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(BASE, "Tugas Copywriting ARE-C Jejak Rimba.docx")

doc = Document()

sec = doc.sections[0]
sec.top_margin = Cm(3)
sec.bottom_margin = Cm(3)
sec.left_margin = Cm(3)
sec.right_margin = Cm(4)


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
    inst = OxmlElement("w:instrText"); inst.set(qn("xml:space"), "preserve"); inst.text = "PAGE"
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1); run._r.append(inst); run._r.append(fld2)
    run.font.name = "Times New Roman"; run.font.size = Pt(11)


add_page_number(sec)

normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
normal.paragraph_format.line_spacing = 1.5
normal.paragraph_format.space_after = Pt(0)


def para(text, align=None, size=12, bold=False, space_before=None, space_after=None, first_line=Cm(1.25)):
    p = doc.add_paragraph()
    p.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r.font.bold = bold
    if p.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
        p.paragraph_format.first_line_indent = first_line
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def center(text, size=12, bold=False, space_after=0):
    return para(text, align=WD_ALIGN_PARAGRAPH.CENTER, size=size, bold=bold,
                space_after=space_after, first_line=Cm(0))


def add_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        r = c.paragraphs[0].add_run(h)
        r.font.bold = True; r.font.size = Pt(12); r.font.name = "Times New Roman"
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(val))
            r.font.size = Pt(12); r.font.name = "Times New Roman"
    return t


# ── Identitas ──
para("Nama : Edgard Mahardika", first_line=Cm(0))
para("Kelas : XI RPL A", first_line=Cm(0))
para("Absen : 10", first_line=Cm(0))
para("", space_after=6, first_line=Cm(0))

# ── Judul ──
center("Tugas Copywriting Landing Page", size=14, bold=True, space_after=2)
center("Matriks ARE-C", size=14, bold=True, space_after=12)

# ── Isi ──
para("Copywriting halaman utama disusun dengan kerangka Assertion, Reason, Evidence, Conclusion (ARE-C). Kerangka ini menata kalimat promosi menjadi alur yang runtut, mulai dari pernyataan, alasan, bukti, hingga ajakan. Klaim yang dipakai hanya mengacu pada fitur yang benar ada di dalam sistem, sehingga isi promosi dapat dibuktikan langsung pada produk.")

add_table(
    ["Elemen", "Kalimat"],
    [
        ["Klaim", "Jejak Rimba, solusi mendaki tanpa beli peralatan mahal."],
        ["Alasan", "Kondisi dan lokasi alat tampil jelas, booking otomatis lewat pilihan tanggal."],
        ["Bukti", "Lokasi penyedia terekam geolokasi, stok diverifikasi sebelum booking."],
        ["Simpulan", "Sewa sekarang, mulai petualanganmu."],
    ],
)
para("Tabel 1. Matriks ARE-C", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, space_before=4)

para("Bukti pada tabel di atas disusun dari fitur yang nyata pada sistem, yaitu pencatatan lokasi penyedia lewat geolokasi serta pemeriksaan stok sebelum pemesanan diterima. Dengan begitu, pesan promosi tidak menjadi klaim kosong, melainkan sesuai dengan cara kerja sistem yang sebenarnya.")

doc.save(OUT)
print("OK ->", OUT)
